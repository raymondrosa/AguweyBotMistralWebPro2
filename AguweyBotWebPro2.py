# ============================================
# AGUWEYBOT - VERSIÓN MINISTRAL-3 (API CORREGIDA)
# CON GUARDADO DE CONVERSACIONES
# ============================================

import os
import base64
import time
import streamlit as st
import streamlit.components.v1 as components
import re
import io
import json
from typing import Optional, Tuple, List, Dict, Any
from datetime import datetime

# CAMBIO: Importar correctamente para Mistral AI (versión actual)
from mistralai import Mistral

# Para documentos
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import chardet

# ============================================
# TEXTO A VOZ
# ============================================
try:
    from gtts import gTTS
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False

# ============================================
# CONFIGURACIÓN
# ============================================
MODEL_NAME = "ministral-3b-latest"  # Modelo de Mistral AI

# Verificar API key
if "MISTRAL_API_KEY" not in st.secrets:
    st.error("❌ No se encontró la API Key de MISTRAL AI")
    st.stop()

# Inicializar cliente Mistral (versión corregida)
client = Mistral(api_key=st.secrets["MISTRAL_API_KEY"])

# Directorio para guardar conversaciones
SAVE_DIR = "conversaciones_guardadas"
os.makedirs(SAVE_DIR, exist_ok=True)

# ============================================
# CONSTANTES Y CONFIGURACIÓN VISUAL
# ============================================
class Config:
    # Colores
    PRIMARY_COLOR = "#00ffff"
    SECONDARY_COLOR = "#00cccc"
    BACKGROUND_DARK = "#0a0c10"
    CARD_BACKGROUND = "#1e2a3a"
    
    # Rutas de archivos
    LOGO_PATH = "logo.png"
    BACKGROUND_PATH = "fondo.png"
    
    # Límites
    MAX_HISTORY_MESSAGES = 10
    MAX_FILE_SIZE_MB = 50
    MAX_CONTEXT_TOKENS = 8000  # Límite para ministral-3

# ============================================
# SYSTEM PROMPT
# ============================================
SYSTEM_PROMPT = """
Eres AguweyBot, un asistente experto en análisis de documentos usando el modelo ministral-3.

Cuando el usuario suba un archivo, DEBES:
1. Leer TODO el contenido del archivo cuidadosamente
2. Responder preguntas específicas sobre su contenido
3. Si te piden resumir, haz un resumen detallado de TODO el documento
4. Si hay datos numéricos, analízalos completamente
5. Si hay código, explícalo línea por línea

REGLAS:
- Usa TODO el contenido del archivo para responder
- No inventes información
- Si no encuentras algo en el archivo, dilo honestamente
- Usa emojis para hacer las respuestas más amigables
- Responde de manera clara, concisa y profesional
- Mantén un tono amigable pero formal
- Sé eficiente en las respuestas (modelo optimizado)
"""

# ============================================
# GESTIÓN DE CONVERSACIONES GUARDADAS
# ============================================
class ConversacionGuardada:
    """Clase para manejar conversaciones guardadas"""
    
    @staticmethod
    def guardar_conversacion(messages: List[Dict], nombre: str = None) -> str:
        """Guarda una conversación en archivo JSON"""
        if not nombre:
            # Generar nombre automático basado en la primera pregunta
            first_user_msg = next((m["content"] for m in messages if m["role"] == "user"), "Nueva conversación")
            nombre = first_user_msg[:50].replace(" ", "_").replace("/", "_").replace("\\", "_")
            nombre = re.sub(r'[^\w\-_\.]', '', nombre)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{SAVE_DIR}/{timestamp}_{nombre}.json"
        
        data = {
            "timestamp": timestamp,
            "nombre": nombre,
            "mensajes": messages,
            "total_mensajes": len(messages),
            "modelo": MODEL_NAME
        }
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return filename
    
    @staticmethod
    def cargar_conversacion(filename: str) -> Optional[List[Dict]]:
        """Carga una conversación desde archivo JSON"""
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data.get("mensajes", [])
        except Exception as e:
            st.error(f"Error al cargar conversación: {str(e)}")
            return None
    
    @staticmethod
    def listar_conversaciones() -> List[Dict[str, Any]]:
        """Lista todas las conversaciones guardadas"""
        conversaciones = []
        
        if not os.path.exists(SAVE_DIR):
            return conversaciones
        
        for filename in os.listdir(SAVE_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(SAVE_DIR, filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    conversaciones.append({
                        "filename": filepath,
                        "nombre": data.get("nombre", "Sin nombre"),
                        "timestamp": data.get("timestamp", "Desconocido"),
                        "total_mensajes": data.get("total_mensajes", 0),
                        "modelo": data.get("modelo", "Desconocido")
                    })
                except:
                    continue
        
        # Ordenar por timestamp (más reciente primero)
        conversaciones.sort(key=lambda x: x["timestamp"], reverse=True)
        return conversaciones
    
    @staticmethod
    def eliminar_conversacion(filename: str) -> bool:
        """Elimina una conversación guardada"""
        try:
            if os.path.exists(filename):
                os.remove(filename)
                return True
        except Exception as e:
            st.error(f"Error al eliminar: {str(e)}")
        return False

# ============================================
# FUNCIÓN PARA TRUNCAR CONTEXTO
# ============================================
def truncar_contexto(texto: str, max_caracteres: int = 6000) -> str:
    """Trunca el contexto para no exceder límites de ministral-3"""
    if len(texto) <= max_caracteres:
        return texto
    
    lines = texto.split('\n')
    result = []
    current_len = 0
    
    for line in lines:
        if current_len + len(line) + 1 <= max_caracteres:
            result.append(line)
            current_len += len(line) + 1
        else:
            remaining = max_caracteres - current_len
            if remaining > 50:
                result.append(line[:remaining] + "...")
            break
    
    return '\n'.join(result)

# ============================================
# FUNCIÓN PARA FONDO
# ============================================
def set_background():
    """Aplica la imagen de fondo si existe con manejo de errores"""
    if os.path.exists(Config.BACKGROUND_PATH):
        try:
            with open(Config.BACKGROUND_PATH, "rb") as f:
                img_data = f.read()
            encoded = base64.b64encode(img_data).decode()
            
            st.markdown(
                f"""
                <style>
                .stApp {{
                    background-image: url("data:image/png;base64,{encoded}");
                    background-size: cover;
                    background-position: center;
                    background-attachment: fixed;
                    background-repeat: no-repeat;
                }}
                
                .main .block-container {{
                    background-color: rgba(0, 0, 0, 0.7);
                    backdrop-filter: blur(10px);
                    border-radius: 20px;
                    padding: 2rem;
                    margin: 2rem auto;
                    border: 1px solid {Config.PRIMARY_COLOR};
                    box-shadow: 0 0 20px rgba(0, 255, 255, 0.3);
                    max-width: 1200px;
                }}
                </style>
                """,
                unsafe_allow_html=True
            )
        except Exception as e:
            st.markdown(f"""
            <style>
            .stApp {{
                background: linear-gradient(135deg, {Config.BACKGROUND_DARK}, #1a1f2a);
            }}
            </style>
            """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <style>
        .stApp {{
            background: linear-gradient(135deg, {Config.BACKGROUND_DARK}, #1a1f2a);
        }}
        </style>
        """, unsafe_allow_html=True)

# ============================================
# ESTILOS CSS (ACTUALIZADOS)
# ============================================
def aplicar_estilos():
    st.markdown(f"""
    <style>
    /* Estilos generales */
    .stApp {{
        background-color: {Config.BACKGROUND_DARK};
    }}
    
    /* Contenedor principal */
    .main .block-container {{
        background-color: rgba(10, 12, 16, 0.85);
        backdrop-filter: blur(10px);
        border-radius: 20px;
        padding: 2rem;
        margin: 1rem auto;
        border: 1px solid {Config.PRIMARY_COLOR};
        box-shadow: 0 0 30px rgba(0, 255, 255, 0.2);
        max-width: 1000px !important;
    }}
    
    /* Títulos */
    h1 {{
        color: {Config.PRIMARY_COLOR} !important;
        font-size: 2.5rem !important;
        text-align: center;
        text-shadow: 0 0 20px rgba(0, 255, 255, 0.5);
        margin-bottom: 0.5rem !important;
        font-weight: bold;
    }}
    
    .subtitle {{
        text-align: center;
        color: #e0e5f0;
        margin-bottom: 2rem;
        font-size: 1.1rem;
    }}
    
    h2, h3 {{
        color: {Config.PRIMARY_COLOR} !important;
    }}
    
    /* Respuestas del asistente */
    .respuesta-aguwey {{
        background: linear-gradient(145deg, {Config.CARD_BACKGROUND}, #15232e);
        border-left: 6px solid {Config.PRIMARY_COLOR};
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        color: white;
        font-size: 1.1rem;
        line-height: 1.6;
        box-shadow: 0 4px 15px rgba(0, 255, 255, 0.1);
    }}
    
    /* Sidebar */
    [data-testid="stSidebar"] {{
        background: linear-gradient(165deg, #0e1219, #0a0e14);
        border-right: 2px solid {Config.PRIMARY_COLOR};
        padding: 1rem;
    }}
    
    /* Botones generales */
    .stButton > button {{
        background: linear-gradient(145deg, {Config.SECONDARY_COLOR}, {Config.PRIMARY_COLOR});
        color: black !important;
        font-weight: bold;
        border: none;
        border-radius: 20px;
        padding: 0.3rem 1rem;
        transition: all 0.2s;
    }}
    
    .stButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0, 255, 255, 0.3);
    }}
    
    /* Botón de copiar */
    .copy-btn {{
        background: rgba(0, 255, 255, 0.1);
        border: 1px solid {Config.PRIMARY_COLOR};
        color: {Config.PRIMARY_COLOR};
        border-radius: 8px;
        padding: 4px 12px;
        cursor: pointer;
        font-size: 12px;
        font-family: sans-serif;
        transition: all 0.3s ease;
        margin-left: 8px;
    }}
    .copy-btn:hover {{ 
        background: {Config.PRIMARY_COLOR}; 
        color: #000;
        transform: translateY(-1px);
        box-shadow: 0 2px 8px rgba(0, 255, 255, 0.3);
    }}
    
    /* Botones de guardar/exportar */
    .save-btn {{
        background: rgba(0, 255, 100, 0.1);
        border: 1px solid #00ff64;
        color: #00ff64;
        border-radius: 8px;
        padding: 4px 12px;
        cursor: pointer;
        font-size: 12px;
        font-family: sans-serif;
        transition: all 0.3s ease;
    }}
    .save-btn:hover {{ 
        background: #00ff64; 
        color: #000;
        transform: translateY(-1px);
    }}
    
    /* Footer */
    .fixed-footer {{
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: rgba(10, 12, 16, 0.98);
        backdrop-filter: blur(12px);
        border-top: 2px solid {Config.PRIMARY_COLOR};
        padding: 0.8rem;
        text-align: center;
        color: #e0e5f0;
        z-index: 999;
        font-size: 0.95rem;
    }}
    
    .fixed-footer strong {{
        color: {Config.PRIMARY_COLOR};
    }}
    
    /* Model badge */
    .model-badge {{
        background: rgba(0, 255, 255, 0.1);
        border: 1px solid {Config.PRIMARY_COLOR};
        border-radius: 20px;
        padding: 2px 8px;
        font-size: 10px;
        display: inline-block;
        margin-left: 10px;
    }}
    
    /* Chat input */
    .stChatInput input {{
        border-radius: 20px;
        border: 1px solid {Config.PRIMARY_COLOR};
        background: rgba(255, 255, 255, 0.05);
        color: white;
    }}
    
    /* Conversaciones guardadas */
    .conversacion-item {{
        background: rgba(0, 255, 255, 0.05);
        border: 1px solid rgba(0, 255, 255, 0.2);
        border-radius: 8px;
        padding: 8px;
        margin-bottom: 8px;
        cursor: pointer;
        transition: all 0.2s;
    }}
    .conversacion-item:hover {{
        background: rgba(0, 255, 255, 0.1);
        border-color: {Config.PRIMARY_COLOR};
    }}
    </style>
    """, unsafe_allow_html=True)

# ============================================
# BOTÓN DE COPIAR
# ============================================
def boton_copiar(texto: str, id_unico: str) -> None:
    """Genera un botón de copiado"""
    texto_escapado = (texto.replace('\\', '\\\\')
                           .replace('`', '\\`')
                           .replace('$', '\\$')
                           .replace('\n', '\\n')
                           .replace("'", "\\'")
                           .replace('"', '\\"'))
    
    html_code = f"""
    <div style="text-align: right; margin-top: 0px;">
        <button id="btn_{id_unico}" class="copy-btn" onclick="copyText_{id_unico}()">
            📋 Copiar
        </button>
    </div>
    <script>
    function copyText_{id_unico}() {{
        const textToCopy = `{texto_escapado}`;
        
        if (navigator.clipboard && navigator.clipboard.writeText) {{
            navigator.clipboard.writeText(textToCopy).then(() => {{
                showCopied_{id_unico}();
            }}).catch(() => {{
                fallbackCopy_{id_unico}(textToCopy);
            }});
        }} else {{
            fallbackCopy_{id_unico}(textToCopy);
        }}
    }}
    
    function fallbackCopy_{id_unico}(text) {{
        const tempTextArea = document.createElement('textarea');
        tempTextArea.value = text;
        tempTextArea.style.position = 'fixed';
        tempTextArea.style.opacity = '0';
        document.body.appendChild(tempTextArea);
        tempTextArea.select();
        document.execCommand('copy');
        document.body.removeChild(tempTextArea);
        showCopied_{id_unico}();
    }}
    
    function showCopied_{id_unico}() {{
        const btn = document.getElementById("btn_{id_unico}");
        const originalText = btn.innerText;
        btn.innerText = "✅ ¡Copiado!";
        btn.style.background = "rgba(0, 255, 0, 0.2)";
        btn.style.borderColor = "#00ff00";
        btn.style.color = "#00ff00";
        
        setTimeout(() => {{ 
            btn.innerText = originalText;
            btn.style.background = "rgba(0, 255, 255, 0.1)";
            btn.style.borderColor = "#00ffff";
            btn.style.color = "#00ffff";
        }}, 2000);
    }}
    </script>
    """
    components.html(html_code, height=40)

# ============================================
# BOTÓN DE GUARDAR
# ============================================
def boton_guardar_conversacion(messages: List[Dict], id_unico: str) -> None:
    """Genera un botón para guardar conversación"""
    html_code = f"""
    <div style="text-align: right; margin-top: 5px;">
        <button id="save_{id_unico}" class="save-btn" onclick="saveConversation_{id_unico}()">
            💾 Guardar conversación
        </button>
    </div>
    <script>
    function saveConversation_{id_unico}() {{
        const btn = document.getElementById("save_{id_unico}");
        const originalText = btn.innerText;
        btn.innerText = "⏳ Guardando...";
        btn.style.background = "rgba(255, 165, 0, 0.2)";
        btn.style.borderColor = "#ffa500";
        btn.style.color = "#ffa500";
        
        // Simular guardado
        setTimeout(() => {{ 
            btn.innerText = "✅ ¡Guardado!";
            btn.style.background = "rgba(0, 255, 0, 0.2)";
            btn.style.borderColor = "#00ff00";
            btn.style.color = "#00ff00";
            
            setTimeout(() => {{ 
                btn.innerText = originalText;
                btn.style.background = "rgba(0, 255, 100, 0.1)";
                btn.style.borderColor = "#00ff64";
                btn.style.color = "#00ff64";
            }}, 2000);
        }}, 1000);
    }}
    </script>
    """
    components.html(html_code, height=40)

# ============================================
# CLASE DATOS ARCHIVO
# ============================================
class DatosArchivo:
    def __init__(self):
        self.nombre: str = ""
        self.contenido_completo: str = ""
        self.tipo: str = ""
        self.dataframe: Optional[pd.DataFrame] = None
        self.num_paginas: int = 0
        self.num_caracteres: int = 0
        self.resumen: str = ""
        self.fecha_carga: float = time.time()
    
    def generar_resumen(self) -> str:
        """Genera un resumen básico del archivo"""
        if self.tipo == "pdf":
            return f"📄 PDF con {self.num_paginas} páginas"
        elif self.tipo in ["excel", "csv"]:
            if self.dataframe is not None:
                return f"📊 Tabla con {len(self.dataframe)} filas y {len(self.dataframe.columns)} columnas"
        elif self.tipo in ["txt", "docx"]:
            palabras = len(self.contenido_completo.split())
            return f"📝 Documento con {palabras} palabras"
        return "📁 Archivo procesado"
    
    def to_dict(self) -> Dict:
        """Convierte a diccionario para guardar"""
        return {
            "nombre": self.nombre,
            "contenido_completo": self.contenido_completo,
            "tipo": self.tipo,
            "num_paginas": self.num_paginas,
            "num_caracteres": self.num_caracteres,
            "resumen": self.resumen,
            "fecha_carga": self.fecha_carga
        }
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'DatosArchivo':
        """Crea desde diccionario"""
        datos = cls()
        datos.nombre = data.get("nombre", "")
        datos.contenido_completo = data.get("contenido_completo", "")
        datos.tipo = data.get("tipo", "")
        datos.num_paginas = data.get("num_paginas", 0)
        datos.num_caracteres = data.get("num_caracteres", 0)
        datos.resumen = data.get("resumen", "")
        datos.fecha_carga = data.get("fecha_carga", time.time())
        return datos

# ============================================
# FUNCIÓN PARA LEER ARCHIVOS
# ============================================
def leer_archivo_completo(uploaded_file):
    """Lee el archivo COMPLETO"""
    if uploaded_file is None:
        return None, "No hay archivo para procesar"
    
    try:
        uploaded_file.seek(0)
        
        uploaded_file.seek(0, os.SEEK_END)
        file_size = uploaded_file.tell()
        uploaded_file.seek(0)
        
        if file_size > Config.MAX_FILE_SIZE_MB * 1024 * 1024:
            return None, f"El archivo excede el límite de {Config.MAX_FILE_SIZE_MB}MB"
        
        nombre = uploaded_file.name.lower()
        datos = DatosArchivo()
        datos.nombre = uploaded_file.name
        
        # PDF
        if nombre.endswith(".pdf"):
            try:
                reader = PdfReader(uploaded_file)
                datos.num_paginas = len(reader.pages)
                texto_completo = []
                
                progress_bar = st.progress(0, text="📖 Leyendo PDF...")
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        texto_completo.append(f"--- PÁGINA {i+1} ---\n{page_text}")
                    progress_bar.progress((i + 1) / datos.num_paginas)
                
                datos.contenido_completo = "\n\n".join(texto_completo)
                datos.tipo = "pdf"
                progress_bar.empty()
                
                if not datos.contenido_completo:
                    return None, "El PDF no contiene texto extraíble"
                    
            except Exception as e:
                return None, f"Error al leer PDF: {str(e)}"
        
        # Excel
        elif nombre.endswith((".xlsx", ".xls")):
            try:
                df = pd.read_excel(uploaded_file)
                datos.dataframe = df
                datos.contenido_completo = f"📊 ARCHIVO EXCEL: {uploaded_file.name}\n"
                datos.contenido_completo += f"Filas: {len(df)}, Columnas: {len(df.columns)}\n"
                datos.contenido_completo += f"Columnas: {', '.join(df.columns.tolist())}\n\n"
                datos.contenido_completo += "DATOS COMPLETOS:\n"
                datos.contenido_completo += df.to_string()
                datos.tipo = "excel"
                
            except Exception as e:
                return None, f"Error al leer Excel: {str(e)}"
        
        # CSV
        elif nombre.endswith(".csv"):
            try:
                raw_data = uploaded_file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                
                df = pd.read_csv(io.BytesIO(raw_data), encoding=encoding)
                datos.dataframe = df
                datos.contenido_completo = f"📊 ARCHIVO CSV: {uploaded_file.name}\n"
                datos.contenido_completo += f"Filas: {len(df)}, Columnas: {len(df.columns)}\n"
                datos.contenido_completo += f"Columnas: {', '.join(df.columns.tolist())}\n\n"
                datos.contenido_completo += "DATOS COMPLETOS:\n"
                datos.contenido_completo += df.to_string()
                datos.tipo = "csv"
                
            except Exception as e:
                return None, f"Error al leer CSV: {str(e)}"
        
        # TXT
        elif nombre.endswith(".txt"):
            try:
                contenido = uploaded_file.read()
                result = chardet.detect(contenido)
                encoding = result['encoding'] or 'utf-8'
                datos.contenido_completo = contenido.decode(encoding)
                datos.tipo = "txt"
                
            except Exception as e:
                return None, f"Error al leer TXT: {str(e)}"
        
        # Word
        elif nombre.endswith(".docx"):
            try:
                doc = Document(uploaded_file)
                texto_completo = []
                
                for p in doc.paragraphs:
                    if p.text.strip():
                        texto_completo.append(p.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        texto_completo.append(" | ".join(row_text))
                
                datos.contenido_completo = "\n".join(texto_completo)
                datos.tipo = "docx"
                
                if not datos.contenido_completo:
                    return None, "El documento no contiene texto"
                    
            except Exception as e:
                return None, f"Error al leer DOCX: {str(e)}"
        
        else:
            return None, f"Tipo de archivo no soportado: {nombre.split('.')[-1]}"
        
        datos.num_caracteres = len(datos.contenido_completo)
        datos.resumen = datos.generar_resumen()
        
        return datos, None
        
    except Exception as e:
        return None, f"Error inesperado: {str(e)}"

# ============================================
# FUNCIÓN PARA STREAMING CON MINISTRAL (CORREGIDA)
# ============================================
def generar_respuesta_streaming(messages, container):
    """Genera respuesta con streaming usando Mistral AI"""
    try:
        full_response = ""
        response_container = container.empty()
        
        # Realizar la solicitud con streaming (API corregida)
        stream_response = client.chat.stream(
            model=MODEL_NAME,
            messages=messages,
            temperature=0.2,
            max_tokens=2000,
        )
        
        start_time = time.time()
        
        # Procesar el stream
        for chunk in stream_response:
            if chunk.data.choices[0].delta.content is not None:
                content = chunk.data.choices[0].delta.content
                full_response += content
                
                elapsed = time.time() - start_time
                response_container.markdown(
                    f'<div class="respuesta-aguwey" style="position: relative;">{full_response}▌<div style="position: absolute; bottom: 5px; right: 10px; font-size: 10px; color: #666;">Generando... {elapsed:.1f}s</div></div>',
                    unsafe_allow_html=True
                )
                time.sleep(0.002)
        
        elapsed = time.time() - start_time
        response_container.markdown(
            f'<div class="respuesta-aguwey" style="position: relative;">{full_response}<div style="position: absolute; bottom: 5px; right: 10px; font-size: 10px; color: #666;">Generado en {elapsed:.1f}s</div></div>',
            unsafe_allow_html=True
        )
        
        return full_response
        
    except Exception as e:
        st.error(f"❌ Error en streaming: {str(e)}")
        return f"Error: {str(e)}"

# ============================================
# TEXTO A VOZ
# ============================================
def texto_a_audio_unico(texto: str) -> Optional[bytes]:
    """Convierte texto a audio"""
    if not TTS_AVAILABLE or not texto or not texto.strip():
        return None
    
    try:
        texto_limpio = re.sub(r'[#*_`\[\]()📄📊🔊🔗🔘🎯✅❌⚠️]', '', texto)
        texto_limpio = re.sub(r'\s+', ' ', texto_limpio).strip()
        
        if not texto_limpio:
            return None
            
        tts = gTTS(text=texto_limpio, lang='es', slow=False)
        
        audio_bytes_io = io.BytesIO()
        tts.write_to_fp(audio_bytes_io)
        audio_bytes_io.seek(0)
        
        return audio_bytes_io.getvalue()
        
    except Exception as e:
        st.warning(f"⚠️ No se pudo generar audio: {str(e)}")
        return None

# ============================================
# MOSTRAR LOGO
# ============================================
def mostrar_logo():
    if os.path.exists(Config.LOGO_PATH):
        try:
            from PIL import Image
            logo = Image.open(Config.LOGO_PATH)
            st.sidebar.image(logo, width=200)
        except:
            st.sidebar.markdown("# 🤖 AguweyBot")
    else:
        st.sidebar.markdown("""
        # 🤖 AguweyBot
        ### *Asistente con Ministral-3*
        """)

def mostrar_info_archivo(datos: DatosArchivo) -> None:
    """Muestra información del archivo cargado"""
    if datos:
        with st.sidebar.expander("📁 Archivo activo", expanded=True):
            st.markdown(f"""
            **Nombre:** {datos.nombre}
            **Tipo:** {datos.resumen}
            **Tamaño:** {datos.num_caracteres:,} caracteres
            **Cargado:** {datetime.fromtimestamp(datos.fecha_carga).strftime('%H:%M:%S')}
            """)
            
            if datos.tipo in ["excel", "csv"] and datos.dataframe is not None:
                st.dataframe(datos.dataframe.head(5), use_container_width=True)
            elif datos.num_caracteres > 500:
                with st.expander("📄 Vista previa"):
                    st.text(datos.contenido_completo[:500] + "...")

# ============================================
# FUNCIÓN PARA EXPORTAR CONVERSACIÓN
# ============================================
def exportar_conversacion(messages: List[Dict]) -> str:
    """Exporta conversación a formato de texto"""
    export_text = "=" * 60 + "\n"
    export_text += f"CONVERSACIÓN AGUWEYBOT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    export_text += "=" * 60 + "\n\n"
    
    for i, msg in enumerate(messages, 1):
        role = "👤 USUARIO" if msg["role"] == "user" else "🤖 AGUWEYBOT"
        export_text += f"[{i}] {role}\n"
        export_text += "-" * 40 + "\n"
        export_text += msg["content"] + "\n"
        export_text += "-" * 40 + "\n\n"
    
    export_text += "=" * 60 + "\n"
    export_text += "FIN DE LA CONVERSACIÓN\n"
    export_text += "=" * 60
    
    return export_text

# ============================================
# FUNCIÓN PRINCIPAL (ACTUALIZADA)
# ============================================
def main():
    st.set_page_config(
        page_title="AguweyBot - Asistente con Ministral-3",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    set_background()
    aplicar_estilos()
    
    # Inicializar session state
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "datos_archivo" not in st.session_state:
        st.session_state.datos_archivo = None
    if "primer_mensaje" not in st.session_state:
        st.session_state.primer_mensaje = True
    if "audio_actual_bytes" not in st.session_state:
        st.session_state.audio_actual_bytes = None
    if "ultimo_audio_idx" not in st.session_state:
        st.session_state.ultimo_audio_idx = -1
    if "conversacion_guardada" not in st.session_state:
        st.session_state.conversacion_guardada = False
    
    # Sidebar
    with st.sidebar:
        mostrar_logo()
        st.markdown("---")
        
        st.markdown("### 🔑 Estado")
        st.success("✅ Mistral AI conectado")
        st.markdown(f"<span style='font-size:12px'>🤖 Modelo: <strong>{MODEL_NAME}</strong></span>", unsafe_allow_html=True)
        if TTS_AVAILABLE:
            st.success("✅ Audio disponible")
        else:
            st.warning("⚠️ Audio no disponible")
        
        st.markdown("---")
        
        # ===== NUEVA SECCIÓN: CONVERSACIONES GUARDADAS =====
        st.markdown("### 💾 Conversaciones Guardadas")
        
        # Botón para guardar conversación actual
        if st.session_state.messages:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Guardar", use_container_width=True, help="Guardar conversación actual"):
                    nombre_predeterminado = None
                    filename = ConversacionGuardada.guardar_conversacion(
                        st.session_state.messages, 
                        nombre_predeterminado
                    )
                    st.session_state.conversacion_guardada = True
                    st.success(f"✅ ¡Conversación guardada!")
                    st.rerun()
            
            with col2:
                # Exportar como texto
                export_text = exportar_conversacion(st.session_state.messages)
                st.download_button(
                    label="📄 Exportar TXT",
                    data=export_text,
                    file_name=f"conversacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        
        st.markdown("---")
        
        # Lista de conversaciones guardadas
        conversaciones = ConversacionGuardada.listar_conversaciones()
        
        if conversaciones:
            st.markdown("**📚 Conversaciones guardadas:**")
            
            for i, conv in enumerate(conversaciones[:5]):  # Mostrar solo las 5 más recientes
                timestamp = conv["timestamp"]
                try:
                    fecha = datetime.strptime(timestamp, "%Y%m%d_%H%M%S").strftime("%d/%m/%Y %H:%M")
                except:
                    fecha = timestamp
                
                col1, col2, col3 = st.columns([6, 1, 1])
                
                with col1:
                    if st.button(f"📝 {conv['nombre'][:30]}... ({fecha})", key=f"load_{i}", use_container_width=True):
                        mensajes_cargados = ConversacionGuardada.cargar_conversacion(conv["filename"])
                        if mensajes_cargados:
                            st.session_state.messages = mensajes_cargados
                            st.success("✅ Conversación cargada")
                            st.rerun()
                
                with col2:
                    # Exportar conversación guardada
                    if st.button("📄", key=f"export_{i}", help="Exportar como TXT"):
                        export_text = exportar_conversacion(conv.get("mensajes", []))
                        st.download_button(
                            label="⬇️",
                            data=export_text,
                            file_name=f"{conv['nombre']}.txt",
                            mime="text/plain",
                            key=f"dl_{i}"
                        )
                
                with col3:
                    if st.button("🗑️", key=f"del_{i}", help="Eliminar"):
                        if ConversacionGuardada.eliminar_conversacion(conv["filename"]):
                            st.success("✅ Eliminada")
                            st.rerun()
        else:
            st.info("📭 No hay conversaciones guardadas")
        
        st.markdown("---")
        
        # ===== SECCIÓN DE ARCHIVOS =====
        st.markdown("### 📎 Subir Archivo")
        
        uploaded_file = st.file_uploader(
            "Elige un archivo",
            type=["pdf", "xlsx", "xls", "csv", "txt", "docx"],
            key="file_uploader",
            label_visibility="collapsed",
            help="Formatos soportados: PDF, Excel, CSV, TXT, DOCX"
        )
        
        if uploaded_file is not None:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📖 Leer TODO", key="btn_leer", use_container_width=True):
                    with st.spinner("📖 Leyendo archivo COMPLETO..."):
                        datos, error = leer_archivo_completo(uploaded_file)
                        
                        if error:
                            st.error(f"❌ {error}")
                        elif datos:
                            st.session_state.datos_archivo = datos
                            st.success(f"✅ {datos.resumen}")
                            st.balloons()
            
            with col2:
                if st.button("🔄 Limpiar", use_container_width=True):
                    st.session_state.datos_archivo = None
                    st.session_state.ultimo_audio_idx = -1
                    st.rerun()
        
        if st.session_state.datos_archivo:
            mostrar_info_archivo(st.session_state.datos_archivo)
        
        st.markdown("---")
        
        if st.button("🔄 Nueva Conversación", use_container_width=True):
            st.session_state.messages = []
            st.session_state.audio_actual_bytes = None
            st.session_state.ultimo_audio_idx = -1
            st.session_state.conversacion_guardada = False
            st.success("¡Conversación reiniciada!")
            st.rerun()
        
        if st.session_state.messages:
            st.markdown("### 📊 Estadísticas")
            user_msgs = sum(1 for m in st.session_state.messages if m["role"] == "user")
            assistant_msgs = sum(1 for m in st.session_state.messages if m["role"] == "assistant")
            st.markdown(f"""
            - 💬 Mensajes: {len(st.session_state.messages)}
            - 👤 Usuario: {user_msgs}
            - 🤖 Asistente: {assistant_msgs}
            """)
    
    # Contenido principal
    st.markdown("<h1>🤖 AguweyBot <span class='model-badge'>Ministral-3</span></h1>", unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Asistente inteligente con análisis de documentos, audio y guardado de conversaciones</p>', unsafe_allow_html=True)
    
    # Mostrar historial
    for i, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            if msg["role"] == "assistant":
                st.markdown(f'<div class="respuesta-aguwey">{msg["content"]}</div>', unsafe_allow_html=True)
                
                col_audio, col_copy, col_spacer = st.columns([1, 1, 4])
                
                with col_audio:
                    if TTS_AVAILABLE:
                        if st.button(f"🔊 Escuchar", key=f"audio_{i}"):
                            with st.spinner("Generando audio..."):
                                audio_bytes = texto_a_audio_unico(msg["content"])
                                if audio_bytes:
                                    st.session_state.audio_actual_bytes = audio_bytes
                                    st.session_state.ultimo_audio_idx = i
                                    st.rerun()
                
                with col_copy:
                    boton_copiar(msg["content"], f"copy_{i}")
                
                if (st.session_state.get('audio_actual_bytes') and 
                    st.session_state.ultimo_audio_idx == i):
                    st.audio(st.session_state.audio_actual_bytes, format="audio/mpeg")
            
            else:
                st.markdown(f"**Tú:** {msg['content']}")
    
    # Mensaje de bienvenida
    if st.session_state.primer_mensaje and not st.session_state.messages:
        st.info("""
        👋 **¡Bienvenido a AguweyBot con Ministral-3!**
        
        **📝 Cómo usar:**
        1. Sube un archivo en el panel izquierdo
        2. Haz clic en **"Leer TODO"**
        3. Pregúntame sobre el contenido
        
        **💾 Nuevo:** ¡Ahora puedes guardar y cargar tus conversaciones!
        - Usa **"Guardar"** para almacenar la conversación actual
        - Usa **"Exportar TXT"** para descargar como archivo de texto
        - Carga conversaciones previas desde la lista en el panel lateral
        
        **🔊 Audio:** Haz clic en "Escuchar" debajo de cualquier respuesta para oírla.
        **📋 Copiar:** Usa el botón "Copiar" para guardar respuestas.
        
        ⚡ **Modelo:** Ministral-3 - Optimizado para respuestas rápidas y eficientes
        """)
        st.session_state.primer_mensaje = False
    
    # Input del usuario
    prompt = st.chat_input("Escribe tu pregunta aquí...", key="chat_input")
    
    if prompt:
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.conversacion_guardada = False  # Marcar que hay cambios sin guardar
        
        with st.chat_message("user"):
            st.markdown(f"**Tú:** {prompt}")
        
        with st.chat_message("assistant"):
            try:
                # Construir mensajes para Mistral (usando diccionarios)
                messages = []
                
                # Agregar sistema
                messages.append({"role": "system", "content": SYSTEM_PROMPT})
                
                # Agregar historial reciente
                ultimos_mensajes = st.session_state.messages[-Config.MAX_HISTORY_MESSAGES:]
                for m in ultimos_mensajes:
                    if m["role"] == "user":
                        messages.append({"role": "user", "content": m["content"]})
                    elif m["role"] == "assistant":
                        messages.append({"role": "assistant", "content": m["content"]})
                
                # Agregar contexto del archivo si existe
                if st.session_state.datos_archivo:
                    datos = st.session_state.datos_archivo
                    
                    contenido_truncado = truncar_contexto(datos.contenido_completo, Config.MAX_CONTEXT_TOKENS)
                    
                    contexto = f"""
📁 ARCHIVO COMPLETO: {datos.nombre}
TIPO: {datos.tipo}
RESUMEN: {datos.resumen}

========== CONTENIDO DEL ARCHIVO ==========

{contenido_truncado}

========== FIN DEL CONTENIDO ==========

PREGUNTA DEL USUARIO: {prompt}

IMPORTANTE: Usa TODO el contenido del archivo para responder.
"""
                    messages.append({"role": "user", "content": contexto})
                else:
                    messages.append({"role": "user", "content": prompt})
                
                # Generar respuesta con streaming
                container = st.empty()
                response = generar_respuesta_streaming(messages, container)
                
                # Guardar respuesta
                st.session_state.messages.append({"role": "assistant", "content": response})
                
                # Limpiar audio anterior
                st.session_state.audio_actual_bytes = None
                st.session_state.ultimo_audio_idx = -1
                
                # Auto-generar audio para respuestas largas
                if TTS_AVAILABLE and len(response) > 100:
                    audio_bytes = texto_a_audio_unico(response)
                    if audio_bytes:
                        st.session_state.audio_actual_bytes = audio_bytes
                        st.session_state.ultimo_audio_idx = len(st.session_state.messages) - 1
                        st.rerun()
                        
            except Exception as e:
                st.error(f"❌ Error al generar respuesta: {str(e)}")
                st.exception(e)
    
    # Footer
    st.markdown(
        f"""
        <div class="fixed-footer">
            <strong>CC-SA</strong> Prof. Raymond Rosa Ávila • AguweyBot con Ministral-3 2026 • 
            <span data-tooltip="Versión optimizada con Mistral AI + Guardado de conversaciones">🚀 v5.0</span>
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()